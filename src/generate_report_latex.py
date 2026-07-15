"""Render LaTeX formulas to PNG and rebuild Markdown / Word reports."""
from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
RES = ROOT / "results"
FIG = ROOT / "figures"
FORMULA_DIR = FIG / "formulas"
REPORT_DIR = ROOT / "report"


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name_cn)


def add_para(doc, text, size=12, bold=False, first_line_indent=True, align="left"):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, name_cn="黑体", size=14 if level == 1 else 12, bold=True)


def add_picture(doc, path: Path, width_cm=14.5, caption: str | None = None):
    if not path.exists():
        add_para(doc, f"[缺少图片: {path.name}]", first_line_indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        set_run_font(run, size=10)


def add_latex_block(doc, latex: str, img_path: Path | None = None, width_cm=12.0):
    """Insert rendered formula image and keep LaTeX source below."""
    if img_path is not None and img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src.add_run(f"LaTeX: ${latex}$" if not latex.strip().startswith(r"\begin") else f"LaTeX: $$ {latex} $$")
    set_run_font(run, name_cn="Consolas", name_en="Consolas", size=9)


def add_table_from_df(doc, df: pd.DataFrame):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.4g}"
            else:
                cells[i].text = "" if pd.isna(val) else str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def render_formula(name: str, latex: str, fontsize: int = 18) -> Path:
    """Render with matplotlib-compatible subset; fall back to simplified string."""
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    path = FORMULA_DIR / f"{name}.png"
    # matplotlib mathtext subset
    tex_mp = (
        latex.replace(r"\dfrac", r"\frac")
        .replace(r"\bigl", "")
        .replace(r"\bigr", "")
        .replace(r"\Bigl", "")
        .replace(r"\Bigr", "")
        .replace(r"\mathrm", r"\mathrm")
        .replace(r"\,", r"\,")
    )
    fig = plt.figure(figsize=(9, 1.35))
    fig.patch.set_facecolor("white")
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    try:
        ax.text(0.5, 0.5, f"${tex_mp}$", fontsize=fontsize, ha="center", va="center")
        fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white", pad_inches=0.15)
    except Exception:
        plt.close(fig)
        fig = plt.figure(figsize=(9, 1.0))
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        ax.text(0.5, 0.5, latex, fontsize=11, ha="center", va="center", family="monospace")
        fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white", pad_inches=0.15)
    plt.close(fig)
    return path


def build_formulas(plant: dict) -> dict[str, tuple[str, Path]]:
    """Return mapping: key -> (full_latex, png_path). Full LaTeX kept for md/tex."""
    specs = {
        "ref_tf": (
            r"G_{\mathrm{ref}}(s)=\frac{0.99\,e^{-190s}}{2895s+1}",
            18,
        ),
        "fopdt": (
            r"G(s)=\frac{K e^{-Ls}}{Ts+1}",
            18,
        ),
        "gain": (
            r"K=\frac{(y(\infty)-y(0))/y_{F.S}}{u/u_{F.S}},\quad y_{F.S}=100,\ u_{F.S}=10",
            13,
        ),
        "k_phys": (
            r"K_{\mathrm{phys}}=K\cdot\frac{y_{F.S}}{u_{F.S}}",
            16,
        ),
        "two_point_classic": (
            r"T=1.5(t_2-t_1),\qquad L=1.5t_1-0.5t_2",
            16,
        ),
        "two_point_log": (
            r"T=\frac{t_2-t_1}{\ln\frac{1-\eta_1}{1-\eta_2}},\qquad L=t_1+T\ln(1-\eta_1)",
            14,
        ),
        "step_response": (
            r"y(t)=y(0)+K_{\mathrm{phys}}u\left(1-e^{-(t-L)/T}\right),\quad t\ge L",
            13,
        ),
        "ident_tf": (
            rf"G_{{\mathrm{{id}}}}(s)=\frac{{{plant['K']:.4f}\,e^{{-{plant['L']:.2f}s}}}}"
            rf"{{{plant['T']:.2f}s+1}}",
            16,
        ),
        "pid": (
            r"u(t)=K_p e(t)+K_i\int_0^{t}e(\tau)\,d\tau+K_d\frac{de(t)}{dt}",
            14,
        ),
        "zn": (
            r"K_p=\frac{1.2T}{K_{\mathrm{phys}}L},\ "
            r"T_i=2L,\ T_d=0.5L,\ "
            r"K_i=\frac{K_p}{T_i},\ K_d=K_p T_d",
            12,
        ),
        "pso_cost": (
            r"J=\alpha\,\mathrm{ISE}+\beta M_p+\gamma t_s+\delta|e_{ss}|+\lambda\overline{(\Delta u)^2}",
            13,
        ),
        "ise": (
            r"\mathrm{ISE}=\int_0^{t_f}(r-y(t))^2\,dt",
            15,
        ),
        "metrics": (
            r"n=\frac{A}{B},\quad M_p=y_{\mathrm{peak}}-r,\quad e_{ss}=y(\infty)-r",
            14,
        ),
        "rmse": (
            r"\mathrm{RMSE}=\sqrt{\frac{1}{N}\sum_{i=1}^{N}(y_i-\hat{y}_i)^2}",
            15,
        ),
        "temp_out": (
            rf"y_{{\mathrm{{out}}}}(t)=\Delta y(t)+T_{{\mathrm{{room}}}},\quad "
            rf"T_{{\mathrm{{room}}}}={plant['T_room']:.2f}^{{\circ}}\mathrm{{C}}",
            13,
        ),
    }
    out = {}
    for key, (latex, fs) in specs.items():
        out[key] = (latex, render_formula(key, latex, fontsize=fs))
    return out


def write_markdown(summary: dict, plant: dict, formulas: dict, df_id: pd.DataFrame, df_m: pd.DataFrame, df_pid: pd.DataFrame):
    mb, ma = summary["metrics_before"], summary["metrics_after"]
    before = df_pid[df_pid["阶段"] == "整定前"].iloc[0]
    after = df_pid[df_pid["阶段"] == "整定后"].iloc[0]
    row = df_id[df_id["方法"] == summary["selected_method"]].iloc[0]
    data = summary.get("data", {})

    md = f"""# 任务2实验报告：控制系统的智能辨识与参数优化

> 姓名：________　学号：________　班级：________

## 1 问题分析

### 1.1 加热炉对象与实验数据分析

加热炉为典型自衡对象。`temperature.csv` 提供开环阶跃响应：时间 $t\\,(\\mathrm{{s}})$，温度 $y\\,(^\\circ\\mathrm{{C}})$，加热电压 $u\\,(\\mathrm{{V}})$。  
数据共 ${data.get('n_samples', 21601)}$ 点，采样周期约 ${data.get('dt', 0.5):.1f}\\,\\mathrm{{s}}$，阶跃电压 $u={data.get('u_step', 3.5)}\\,\\mathrm{{V}}$，  
$y(0)\\approx {data.get('y0', 16.90):.2f}^\\circ\\mathrm{{C}}$，$y(\\infty)\\approx {data.get('yss', 51.27):.2f}^\\circ\\mathrm{{C}}$。

对象采用 **一阶惯性加纯滞后（FOPDT）** 模型：

$$
{formulas['fopdt'][0]}
$$

课程标准参考传递函数（仅作对比，不直接作为控制模型）：

$$
{formulas['ref_tf'][0]}
$$

其中百分比量程 $y_{{F.S}}=100$，$u_{{F.S}}=10$，物理温升增益

$$
{formulas['k_phys'][0]}
$$

闭环时输出叠加室温偏置：$y_{{\\mathrm{{out}}}}(t)=\\Delta y(t)+T_{{\\mathrm{{room}}}}$。

![图1 阶跃响应](../figures/01_step_response_data.png)

### 1.2 辨识、控制与优化方案

1. 多方法辨识 FOPDT 参数 $(K,T,L)$，并以 RMSE / 参数相对误差对照参考模型；  
2. 基于辨识模型设计单位负反馈 PID，设定值 $r=35^\\circ\\mathrm{{C}}$，控制量限幅 $0\\sim 10\\,\\mathrm{{V}}$；  
3. 用 PSO 优化 $(K_p,K_i,K_d)$；  
4. 计算衰减比、最大偏差、5%回复时间与余差。

![图2 闭环结构](../figures/09_control_structure.png)

## 2 实验过程

### 2.1 数据处理与系统辨识

百分比增益：

$$
{formulas['gain'][0]}
$$

阶跃响应解析式：

$$
{formulas['step_response'][0]}
$$

**经典两点法**（$\\eta_1\\approx 0.393,\\ \\eta_2\\approx 0.632$）：

$$
{formulas['two_point_classic'][0]}
$$

**对数两点法**：

$$
{formulas['two_point_log'][0]}
$$

另实现切线法、面积/Smith 法与约束最小二乘。拟合误差：

$$
{formulas['rmse'][0]}
$$

### 2.2 模型准确性验证

本实验采用 **{summary['selected_method']}**，辨识模型为

$$
{formulas['ident_tf'][0]}
$$

相对参考模型： $\\Delta K={row['ΔK%']:.2f}\\%$，$\\Delta T={row['ΔT%']:.2f}\\%$，$\\Delta L={row['ΔL%']:.2f}\\%$，  
参数平均相对误差约 ${row['参数平均相对误差%']:.2f}\\%$；RMSE$={row['RMSE']:.4f}^\\circ\\mathrm{{C}}$，$R^2={row['R2']:.5f}$。

![图3 拟合对比](../figures/02_identification_fit.png)

![图4 残差](../figures/03_residuals.png)

![图5 参数对比](../figures/04b_params_vs_reference.png)

### 2.3 PID 闭环控制设计

并联 PID：

$$
{formulas['pid'][0]}
$$

Ziegler–Nichols 反应曲线整定（作为整定前基线参考）：

$$
{formulas['zn'][0]}
$$

输出方程：

$$
{formulas['temp_out'][0]}
$$

### 2.4 智能优化 PID 参数

PSO 适应度：

$$
{formulas['pso_cost'][0]}
$$

其中

$$
{formulas['ise'][0]}
$$

![图7 PSO收敛](../figures/07_pso_convergence.png)

### 2.5 多种辨识方法对比

| 方法 | $K$ | $T\\,(\\mathrm{{s}})$ | $L\\,(\\mathrm{{s}})$ | RMSE | $\\Delta K\\%$ | $\\Delta T\\%$ | $\\Delta L\\%$ | 参数平均相对误差% |
|------|-----|----------------------|----------------------|------|--------------|--------------|--------------|-------------------|
"""
    for _, r in df_id.iterrows():
        md += (
            f"| {r['方法']} | {r['K']:.4g} | {r['T(s)']:.2f} | {r['L(s)']:.2f} | "
            f"{r['RMSE']:.4f} | {r['ΔK%']:.2f} | {r['ΔT%']:.2f} | {r['ΔL%']:.2f} | "
            f"{r['参数平均相对误差%']:.2f} |\n"
        )

    md += f"""
![图6 RMSE对比](../figures/04_rmse_bar.png)

## 3 实验结果及分析

### 3.1 系统辨识结果

最终模型

$$
{formulas['ident_tf'][0]}
$$

与

$$
{formulas['ref_tf'][0]}
$$

对比，平均参数相对误差约 ${row['参数平均相对误差%']:.2f}\\%$，满足误差不宜过大的要求。

### 3.2 PID 参数优化结果

整定前：
$$
K_p={before['Kp']:.4f},\\quad K_i={before['Ki']:.6f},\\quad K_d={before['Kd']:.4f}
$$

整定后：
$$
K_p={after['Kp']:.4f},\\quad K_i={after['Ki']:.6f},\\quad K_d={after['Kd']:.4f}
$$

![图8 闭环温度](../figures/05_closed_loop_temperature.png)

![图9 控制量](../figures/06_control_signal.png)

### 3.3 动态与稳态指标分析

指标定义：

$$
{formulas['metrics'][0]}
$$

| 阶段 | 衰减比 $n$ | 最大偏差 $M_p\\,(^\\circ\\mathrm{{C}})$ | 5%回复时间 $t_s\\,(\\mathrm{{s}})$ | 余差 $e_{{ss}}\\,(^\\circ\\mathrm{{C}})$ |
|------|-----------|--------------------------------------|-----------------------------------|------------------------------------------|
| 整定前 | {mb['衰减比'] if mb['衰减比'] is not None else '—'} | {mb['最大偏差']:.3f} | {mb['5%回复时间']:.0f} | {mb['余差']:.3f} |
| 整定后 | {'—' if ma.get('衰减比') in (None, float('inf')) or (isinstance(ma.get('衰减比'), float) and (ma['衰减比']!=ma['衰减比'])) else ma['衰减比']} | {ma['最大偏差']:.4f} | {ma['5%回复时间']:.0f} | {ma['余差']:.5f} |

![图10 指标对比](../figures/08_performance_metrics.png)

### 3.4 辨识方法对比

约束最小二乘与对数两点法综合表现最好；切线法对噪声敏感；面积法易低估滞后。

## 4 总结体会

完成了 FOPDT 多方法辨识、与标准传递函数对比、PID 设计与 PSO 智能整定。辨识模型参数误差约 $5\\%$ 量级，整定后温度可平稳跟踪 $35^\\circ\\mathrm{{C}}$。

## 5 附录

- 仓库链接：`https://github.com/<your-username>/furnace-pid-identification`
- 复现：`python src/run_experiment.py`，`python src/generate_report_latex.py`

### 公式一览（LaTeX）

"""
    for key, (latex, _) in formulas.items():
        md += f"- `{key}`:\n\n$$\n{latex}\n$$\n\n"

    out = REPORT_DIR / "任务2_实验报告_LaTeX.md"
    out.write_text(md, encoding="utf-8")
    print("Markdown report:", out)

    # Also a standalone .tex snippet bank
    tex = "% Auto-generated formula bank for Task 2\n"
    for key, (latex, _) in formulas.items():
        tex += f"% {key}\n\\begin{{equation}}\n{latex}\n\\end{{equation}}\n\n"
    tex_path = REPORT_DIR / "formulas.tex"
    tex_path.write_text(tex, encoding="utf-8")
    print("TeX formulas:", tex_path)


def write_docx(summary, plant, formulas, df_id, df_m, df_pid):
    method = summary["selected_method"]
    row = df_id[df_id["方法"] == method].iloc[0]
    before = df_pid[df_pid["阶段"] == "整定前"].iloc[0]
    after = df_pid[df_pid["阶段"] == "整定后"].iloc[0]
    mb, ma = summary["metrics_before"], summary["metrics_after"]
    data = summary.get("data", {})

    doc = Document()
    section = doc.sections[0]
    for m in ("top_margin", "bottom_margin"):
        setattr(section, m, Cm(2.54))
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("任务2实验报告\n控制系统的智能辨识与参数优化")
    set_run_font(r, name_cn="黑体", size=16, bold=True)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(info.add_run("姓名：________    学号：________    班级：________"), size=12)

    add_heading_cn(doc, "1 问题分析", 1)
    add_heading_cn(doc, "1.1 加热炉对象与实验数据分析", 2)
    add_para(
        doc,
        "加热炉是过程控制中的典型自衡对象。本实验使用 temperature.csv 提供的开环阶跃响应数据。"
        f"数据共 {data.get('n_samples', 21601)} 个采样点，采样周期约 {data.get('dt', 0.5):.1f} s，"
        f"阶跃电压 {data.get('u_step', 3.5)} V，y(0)≈{data.get('y0', 16.9):.2f}℃，"
        f"y(∞)≈{data.get('yss', 51.27):.2f}℃。对象采用 FOPDT 模型，公式如下（LaTeX 渲染）。",
    )
    add_latex_block(doc, formulas["fopdt"][0], formulas["fopdt"][1], 10)
    add_para(doc, "课程标准参考传递函数（仅作对比验证）：", first_line_indent=True)
    add_latex_block(doc, formulas["ref_tf"][0], formulas["ref_tf"][1], 11)
    add_para(doc, "百分比增益与物理增益：", first_line_indent=True)
    add_latex_block(doc, formulas["gain"][0], formulas["gain"][1], 13)
    add_latex_block(doc, formulas["k_phys"][0], formulas["k_phys"][1], 9)
    add_picture(doc, FIG / "01_step_response_data.png", caption="图1 加热炉 3.5 V 阶跃响应实测曲线")

    add_heading_cn(doc, "1.2 辨识、控制与优化方案", 2)
    add_para(
        doc,
        "整体方案：多方法辨识 → 模型验证（对照参考传递函数）→ PID 闭环（设定值 35℃）→ PSO 智能整定 → 动/稳态指标分析。",
    )
    add_picture(doc, FIG / "09_control_structure.png", caption="图2 加热炉 PID 闭环控制结构")

    add_heading_cn(doc, "2 实验过程", 1)
    add_heading_cn(doc, "2.1 数据处理与系统辨识", 2)
    add_para(doc, "阶跃响应解析式、经典两点法与对数两点法公式如下。")
    add_latex_block(doc, formulas["step_response"][0], formulas["step_response"][1], 13)
    add_latex_block(doc, formulas["two_point_classic"][0], formulas["two_point_classic"][1], 11)
    add_latex_block(doc, formulas["two_point_log"][0], formulas["two_point_log"][1], 13)
    add_latex_block(doc, formulas["rmse"][0], formulas["rmse"][1], 11)

    add_heading_cn(doc, "2.2 模型准确性验证", 2)
    add_para(doc, f"选取“{method}”，辨识得到：")
    add_latex_block(doc, formulas["ident_tf"][0], formulas["ident_tf"][1], 12)
    add_para(
        doc,
        f"相对参考模型：ΔK={row['ΔK%']:.2f}%，ΔT={row['ΔT%']:.2f}%，ΔL={row['ΔL%']:.2f}%，"
        f"平均相对误差约 {row['参数平均相对误差%']:.2f}%；RMSE={row['RMSE']:.4f}℃，R²={row['R2']:.5f}。",
    )
    add_picture(doc, FIG / "02_identification_fit.png", caption="图3 辨识模型拟合对比")
    add_picture(doc, FIG / "03_residuals.png", caption="图4 残差对比")
    add_picture(doc, FIG / "04b_params_vs_reference.png", caption="图5 参数与参考模型对比")

    add_heading_cn(doc, "2.3 PID 闭环控制设计", 2)
    add_latex_block(doc, formulas["pid"][0], formulas["pid"][1], 12)
    add_latex_block(doc, formulas["zn"][0], formulas["zn"][1], 13)
    add_latex_block(doc, formulas["temp_out"][0], formulas["temp_out"][1], 13)

    add_heading_cn(doc, "2.4 智能优化 PID 参数", 2)
    add_para(doc, "PSO 适应度函数（含 ISE、超调、调节时间、余差与控制抖振惩罚）：")
    add_latex_block(doc, formulas["pso_cost"][0], formulas["pso_cost"][1], 12)
    add_latex_block(doc, formulas["ise"][0], formulas["ise"][1], 10)
    add_picture(doc, FIG / "07_pso_convergence.png", caption="图7 PSO 适应度收敛曲线")

    add_heading_cn(doc, "2.5 多种辨识方法对比", 2)
    add_table_from_df(doc, df_id)
    add_picture(doc, FIG / "04_rmse_bar.png", caption="图6 各辨识方法 RMSE 对比")

    add_heading_cn(doc, "3 实验结果及分析", 1)
    add_heading_cn(doc, "3.1 系统辨识结果", 2)
    add_para(doc, "辨识模型与参考模型对照如下，参数平均相对误差约 "
            f"{row['参数平均相对误差%']:.2f}%，满足误差不宜过大的要求。")
    add_latex_block(doc, formulas["ident_tf"][0], formulas["ident_tf"][1], 11)
    add_latex_block(doc, formulas["ref_tf"][0], formulas["ref_tf"][1], 11)

    add_heading_cn(doc, "3.2 PID 参数优化结果", 2)
    add_para(
        doc,
        f"整定前：Kp={before['Kp']:.4f}，Ki={before['Ki']:.6f}，Kd={before['Kd']:.4f}；"
        f"整定后：Kp={after['Kp']:.4f}，Ki={after['Ki']:.6f}，Kd={after['Kd']:.4f}。",
    )
    add_picture(doc, FIG / "05_closed_loop_temperature.png", caption="图8 闭环温度响应对比")
    add_picture(doc, FIG / "06_control_signal.png", caption="图9 控制量对比")

    add_heading_cn(doc, "3.3 动态与稳态指标分析", 2)
    add_latex_block(doc, formulas["metrics"][0], formulas["metrics"][1], 12)
    add_table_from_df(doc, df_m)
    add_para(
        doc,
        f"整定前：最大偏差 {mb['最大偏差']:.3f}℃，5%回复时间 {mb['5%回复时间']:.0f} s，余差 {mb['余差']:.3f}℃；"
        f"整定后：最大偏差 {ma['最大偏差']:.4f}℃，5%回复时间 {ma['5%回复时间']:.0f} s，余差 {ma['余差']:.5f}℃。",
    )
    add_picture(doc, FIG / "08_performance_metrics.png", caption="图10 动态与稳态指标对比")

    add_heading_cn(doc, "3.4 辨识方法对比", 2)
    add_para(doc, "约束最小二乘与对数两点法综合最优；切线法对噪声敏感；面积法易低估滞后。")

    add_heading_cn(doc, "4 总结体会", 1)
    add_para(
        doc,
        "本实验完成了加热炉 FOPDT 多方法辨识、与课程标准传递函数定量对比、PID 闭环及 PSO 智能整定。"
        "辨识参数平均误差约 5%，整定后可平稳跟踪 35℃。",
    )

    add_heading_cn(doc, "5 附录", 1)
    add_para(doc, "仓库链接：https://github.com/<your-username>/furnace-pid-identification", first_line_indent=False)
    add_para(doc, "Markdown(LaTeX) 报告：report/任务2_实验报告_LaTeX.md；公式库：report/formulas.tex", first_line_indent=False)
    add_para(doc, "复现：python src/run_experiment.py && python src/generate_report_latex.py", first_line_indent=False)

    out = REPORT_DIR / "任务2_控制系统的智能辨识与参数优化_实验报告.docx"
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("Word report:", out)
    return out


def main():
    plt.rcParams["mathtext.fontset"] = "cm"
    summary = json.loads((RES / "summary.json").read_text(encoding="utf-8"))
    plant = summary["identified_model"]
    df_id = pd.read_csv(RES / "identification_comparison.csv")
    df_m = pd.read_csv(RES / "performance_metrics.csv")
    df_pid = pd.read_csv(RES / "pid_parameters.csv")

    formulas = build_formulas(plant)
    write_markdown(summary, plant, formulas, df_id, df_m, df_pid)
    docx_path = write_docx(summary, plant, formulas, df_id, df_m, df_pid)

    # copy to workspace root
    root_copy = ROOT.parent / "任务2_实验报告.docx"
    root_copy.write_bytes(docx_path.read_bytes())
    print("Copied to", root_copy)


if __name__ == "__main__":
    main()
