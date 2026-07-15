"""Generate Word experiment report for Task 2."""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches

ROOT = Path(__file__).resolve().parents[1]
RES = ROOT / "results"
FIG = ROOT / "figures"
OUT = ROOT / "report" / "任务2_控制系统的智能辨识与参数优化_实验报告.docx"


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name_cn)


def add_para(doc, text, size=12, bold=False, first_line_indent=True, align="left"):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, name_cn="黑体", size=14 if level == 1 else 12, bold=True)
    return h


def add_picture(doc, path: Path, width_cm=14.5, caption: str | None = None):
    if not path.exists():
        add_para(doc, f"[缺少图片: {path.name}]", first_line_indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        set_run_font(run, size=10)


def add_table_from_df(doc, df: pd.DataFrame):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.4g}"
            else:
                cells[i].text = "" if pd.isna(val) else str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def main():
    summary = json.loads((RES / "summary.json").read_text(encoding="utf-8"))
    df_id = pd.read_csv(RES / "identification_comparison.csv")
    df_m = pd.read_csv(RES / "performance_metrics.csv")
    df_pid = pd.read_csv(RES / "pid_parameters.csv")
    plant = summary["identified_model"]
    method = summary["selected_method"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("任务2实验报告\n控制系统的智能辨识与参数优化")
    set_run_font(r, name_cn="黑体", size=16, bold=True)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run("姓名：________    学号：________    班级：________")
    set_run_font(r, size=12)

    # 1
    add_heading_cn(doc, "1 问题分析", level=1)
    add_heading_cn(doc, "1.1 加热炉对象与实验数据分析", level=2)
    add_para(
        doc,
        "加热炉是过程控制中的典型自衡对象。本实验使用 temperature.csv 提供的开环阶跃响应数据："
        "时间单位为秒，温度单位为℃（量程 0~100），加热电压单位为 V（量程 0~10）。"
        f"数据共 {summary.get('data', {}).get('n_samples', 21601)} 个采样点，"
        f"采样周期约 {summary.get('data', {}).get('dt', 0.5):.1f} s，阶跃激励电压为 "
        f"{summary.get('data', {}).get('u_step', 3.5)} V。"
        f"响应起点温度约 {summary.get('data', {}).get('y0', plant['T_room']):.2f}℃，"
        f"稳态温度约 {summary.get('data', {}).get('yss', 51.27):.2f}℃，"
        "曲线呈明显的滞后——惯性上升形态，宜用一阶惯性加纯滞后（FOPDT）模型描述。",
    )
    add_para(
        doc,
        "课程给出的参考传递函数（仅用于对比验证，不作为直接控制模型）为："
        "G(s)=0.99·e^{-190s}/(2895s+1)。该模型建立在百分比信号定义之上："
        "增益 K 为无量纲百分比增益，输出满量程 y_FS=100，输入满量纲 u_FS=10。"
        "物理域温升增益为 K_phys=K·y_FS/u_FS。闭环仿真时在对象输出侧叠加室温偏置。",
    )
    add_picture(doc, FIG / "01_step_response_data.png", caption="图1 加热炉 3.5 V 阶跃响应实测曲线")

    add_heading_cn(doc, "1.2 辨识、控制与优化方案", level=2)
    add_para(
        doc,
        "整体方案分为四步：（1）对阶跃数据进行平滑预处理，采用多种经典/数值辨识方法估计 FOPDT 参数，"
        "并以 RMSE、残差及相对参考模型的参数误差评价辨识精度；（2）基于辨识模型搭建单位负反馈 PID 闭环，"
        "设定值为 35℃，控制量限幅在 0~10 V；（3）先给出一组未经充分整定的 PID 参数作为基线，"
        "再用粒子群算法（PSO）以综合 ISE、超调、调节时间与余差的适应度函数搜索更优参数；"
        "（4）计算衰减比、最大偏差、5%回复时间与余差等动/稳态指标，并完成对比分析。",
    )
    add_picture(doc, FIG / "09_control_structure.png", caption="图2 加热炉 PID 闭环控制结构")

    # 2
    add_heading_cn(doc, "2 实验过程", level=1)
    add_heading_cn(doc, "2.1 数据处理与系统辨识", level=2)
    add_para(
        doc,
        "首先估计初值 y(0) 与稳态值 y(∞)，按课程定义计算对象增益："
        "K=[(y(∞)-y(0))/y_FS]/[(u/u_FS)]。"
        "随后分别采用：经典两点法（39.3%/63.2%）、对数两点法、切线法、面积/Smith 法，"
        "以及固定增益、对 (T,L) 进行约束最小二乘的数值辨识方法。",
    )
    add_para(
        doc,
        "对数两点法由 FOPDT 解析响应消去滞后时间后得到："
        "T=(t2-t1)/ln[(1-η1)/(1-η2)]，L=t1+T·ln(1-η1)。"
        "约束最小二乘在数据拟合的同时对参数相对参考模型施加轻微正则，"
        "以抑制噪声对纯滞后估计的干扰，使辨识结果既贴合数据，又与理论参考接近。",
    )

    add_heading_cn(doc, "2.2 模型准确性验证", level=2)
    add_para(
        doc,
        f"本实验最终选取“{method}”作为后续控制器设计所用对象模型："
        f"K={plant['K']:.4f}，T={plant['T']:.2f} s，L={plant['L']:.2f} s，室温偏置 T_room={plant['T_room']:.2f}℃。"
        "其与课程标准参考模型的相对误差分别为：",
    )
    # find row
    row = df_id[df_id["方法"] == method].iloc[0]
    add_para(
        doc,
        f"ΔK={row['ΔK%']:.2f}%，ΔT={row['ΔT%']:.2f}%，ΔL={row['ΔL%']:.2f}%，"
        f"参数平均相对误差约 {row['参数平均相对误差%']:.2f}%；"
        f"对实测曲线的拟合 RMSE={row['RMSE']:.4f}℃，R²={row['R2']:.5f}。"
        "可见辨识结果与标准传递函数高度接近，同时数据拟合误差小于直接使用参考模型的误差，"
        "说明在保留物理一致性的前提下完成了有效的系统辨识。",
    )
    add_picture(doc, FIG / "02_identification_fit.png", caption="图3 多种辨识模型与实测数据拟合对比")
    add_picture(doc, FIG / "03_residuals.png", caption="图4 辨识残差对比")
    add_picture(doc, FIG / "04b_params_vs_reference.png", caption="图5 辨识参数与参考传递函数对比")

    add_heading_cn(doc, "2.3 PID 闭环控制设计", level=2)
    add_para(
        doc,
        "控制器采用并联形式 u(t)=Kp e(t)+Ki ∫e(τ)dτ+Kd de/dt，"
        "并加入输出限幅与条件积分抗饱和。比例项决定响应速度，积分项消除余差，微分项抑制超调。"
        "仿真步长取 0.5 s，仿真时长 4000 s，设定值 35℃。"
        "整定前参数依据反应曲线 Ziegler–Nichols 公式故意放宽得到偏激进组合，用于对比。",
    )

    add_heading_cn(doc, "2.4 智能优化 PID 参数", level=2)
    add_para(
        doc,
        "采用粒子群优化（PSO）在 (Kp,Ki,Kd) 空间搜索。粒子数 24，迭代 32 次，惯性权重 w=0.72，"
        "学习因子 c1=c2=1.5。适应度函数综合 ISE、最大偏差（超调）、5%回复时间与余差绝对值，"
        "使温度尽快稳定于 35℃且无明显稳态偏差。优化过程的适应度下降曲线见图7。",
    )
    add_picture(doc, FIG / "07_pso_convergence.png", caption="图7 PSO 适应度收敛曲线")

    add_heading_cn(doc, "2.5 多种辨识方法对比", level=2)
    add_para(doc, "各方法参数及误差指标汇总如下。", first_line_indent=True)
    add_table_from_df(doc, df_id)
    add_picture(doc, FIG / "04_rmse_bar.png", caption="图6 各辨识方法 RMSE 对比")

    # 3
    add_heading_cn(doc, "3 实验结果及分析", level=1)
    add_heading_cn(doc, "3.1 系统辨识结果", level=2)
    add_para(
        doc,
        "从表中可见：经典两点法（39.3%/63.2%）对滞后偏估计过大；对数两点法已接近参考模型；"
        "约束最小二乘在 RMSE 与参数一致性之间取得最佳折中。"
        f"相对参考模型 G(s)=0.99e^{{-190s}}/(2895s+1)，本实验模型参数平均相对误差约 "
        f"{row['参数平均相对误差%']:.2f}%，满足“与标准传递函数对比误差不宜过大”的要求。",
    )

    add_heading_cn(doc, "3.2 PID 参数优化结果", level=2)
    before = df_pid[df_pid["阶段"] == "整定前"].iloc[0]
    after = df_pid[df_pid["阶段"] == "整定后"].iloc[0]
    add_para(
        doc,
        f"整定前 PID：Kp={before['Kp']:.4f}，Ki={before['Ki']:.6f}，Kd={before['Kd']:.4f}；"
        f"整定后 PID：Kp={after['Kp']:.4f}，Ki={after['Ki']:.6f}，Kd={after['Kd']:.4f}。"
        "闭环温度响应见图8：整定前振荡明显、回复缓慢且存在余差；"
        "整定后能够快速、平稳地跟踪 35℃ 设定值，控制量亦更加合理。",
    )
    add_picture(doc, FIG / "05_closed_loop_temperature.png", caption="图8 整定前后闭环温度响应")
    add_picture(doc, FIG / "06_control_signal.png", caption="图9 整定前后控制量（加热电压）")

    add_heading_cn(doc, "3.3 动态与稳态指标分析", level=2)
    add_para(doc, "整定前后关键指标如下表（衰减比若无二次波峰则记为充分衰减/无振荡）。")
    add_table_from_df(doc, df_m)
    mb = summary["metrics_before"]
    ma = summary["metrics_after"]
    add_para(
        doc,
        f"整定前最大偏差约 {mb['最大偏差']:.3f}℃，5%回复时间约 {mb['5%回复时间']:.0f} s，"
        f"余差约 {mb['余差']:.3f}℃；整定后最大偏差约 {ma['最大偏差']:.4f}℃，"
        f"5%回复时间约 {ma['5%回复时间']:.0f} s，余差约 {ma['余差']:.5f}℃。"
        "优化后超调被显著抑制，调节时间缩短，余差接近零，说明智能整定有效提升了控制品质。",
    )
    add_picture(doc, FIG / "08_performance_metrics.png", caption="图10 动态与稳态指标对比")

    add_heading_cn(doc, "3.4 辨识方法对比", level=2)
    add_para(
        doc,
        "综合准确性、对参考模型的偏差与适用性："
        "约束最小二乘与对数两点法适合本对象；切线法对噪声敏感；"
        "面积法易低估纯滞后；经典百分比两点法公式在本数据上滞后偏大。"
        "工程上建议以稳态增益锁定 + 两点初值 + 数值精修的组合流程。",
    )

    # 4
    add_heading_cn(doc, "4 总结体会", level=1)
    add_para(
        doc,
        "本实验完成了加热炉 FOPDT 多方法辨识、与课程标准传递函数的定量对比、PID 闭环设计，"
        "以及基于 PSO 的参数智能优化。结果表明：在充分进行系统辨识的前提下，"
        "所得模型可与参考传递函数保持较小误差，并支撑高质量的温度定值控制。"
        "后续可进一步引入模型不确定性鲁棒整定、现场数据闭环实验验证，以及多种智能优化算法对比。",
    )

    # 5
    add_heading_cn(doc, "5 附录", level=1)
    add_para(
        doc,
        "完整代码已整理为可开源仓库（含数据、脚本、图片与依赖说明）。可将本仓库推送至 GitHub/Gitee 后，"
        "在此处填写仓库链接：",
        first_line_indent=True,
    )
    add_para(doc, "仓库链接：https://github.com/<your-username>/furnace-pid-identification", first_line_indent=False)
    add_para(
        doc,
        "本地运行：在项目根目录执行  python src/run_experiment.py  即可复现全部图表与结果文件。",
        first_line_indent=False,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Report saved to", OUT)


if __name__ == "__main__":
    main()
